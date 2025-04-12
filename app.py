from flask import Flask, render_template, request, send_file
import pandas as pd
from scholarly import scholarly
import requests
import io
from docx import Document
app = Flask(__name__)

def fetch_author_details(authors_list):

    author_info = []
    for name in authors_list:
        search_query = scholarly.search_author(name)
        author = next(search_query)
        author = scholarly.fill(author)
        author_info.append({
            "name": author.get('name', 'N/A'),
            "email": author.get('email', 'N/A, unable to fetch as email is often masked'),
            "affiliation": author.get('affiliation', 'N/A'),
            "citations": author.get('citedby', 'N/A')
        })
    return author_info
def fetch_publications(author_name, start_year=None, end_year=None):
    try:
        search_query = scholarly.search_author(author_name)
        author = next(search_query)
        scholarly.fill(author, sections=['publications'])
        publications = author['publications']
        
        journals = {}
        conferences = {}
        miscellaneous = []
        total_citations = 0
        total_papers = 0

        for pub in publications:
            pub_info = pub.get('bib', {})
            year = pub_info.get('pub_year', 'Unknown')

            if year == 'Unknown' or not year.isdigit():
                year_label = 'Unknown Year'
                details = {
                    'Title': pub_info.get('title', 'No Title'),
                    'Citation Link': f"https://scholar.google.co.in/citations?view_op=view_citation&hl=en&user={author['scholar_id']}&citation_for_view={pub.get('author_pub_id', 'No Citation ID')}",
                    'Venue': pub_info.get('venue', 'N/A'),
                    'Publisher': pub_info.get('publisher', 'N/A'),
                    'Cited By': pub.get('num_citations', 'N/A'),
                    'Year': year_label,
                }
                miscellaneous.append(details)
                continue
            
            year_label = int(year)

            if start_year is not None and end_year is not None:
                if not (start_year <= year_label <= end_year):
                    continue

            total_citations += int(pub.get('num_citations', 0))
            total_papers += 1

            title = pub_info.get('title', 'No Title')
            citation_id = pub.get('author_pub_id', 'No Citation ID')
            citation_url = f"https://scholar.google.co.in/citations?view_op=view_citation&hl=en&user={author['scholar_id']}&citation_for_view={citation_id}"

            details = {
                'Title': title,
                'Citation Link': citation_url,
                'Venue': pub_info.get('venue', 'N/A'),
                'Publisher': pub_info.get('publisher', 'N/A'),
                'Cited By': pub.get('num_citations', 'N/A'),
                'Year': year_label,
            }

            crossref_data = fetch_crossref_details(title)
            details.update(crossref_data)

            if 'conference' in details.get('journal_or_conference', '').lower():
                conferences.setdefault(year_label, []).append(details)
            elif 'journal' in details.get('journal_or_conference', '').lower():
                journals.setdefault(year_label, []).append(details)
            else:
                miscellaneous.append(details)

        summary_text = (f"Between {start_year} and {end_year}, {author_name} published a total of {total_papers} research papers "
                        f"with a cumulative citation count of {total_citations}.")

        return journals, conferences, miscellaneous, summary_text
    
    except Exception as e:
        print(f"An error occurred: {e}")
        return None, None, None, "Summary not available."


def fetch_crossref_details(title):  
    url = f"https://api.crossref.org/works"
    params = {'query.title': title}
    response = requests.get(url, params=params)

    if response.status_code == 200:
        data = response.json()
        items = data.get('message', {}).get('items', [])
        if items:
            item = items[0]
            return {
                'journal_or_conference': item.get('container-title', ['N/A'])[0],
                'publisher': item.get('publisher', 'N/A')
            }
    return {'journal_or_conference': 'N/A', 'publisher': 'N/A'}

def save_to_excel(journals=None, conferences=None, miscellaneous=None, selected_type='all'):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        
        if selected_type in ['all', 'journal']:
            journal_data = []
            if journals:
                for year, papers in sorted(journals.items()):
                    for paper in papers:
                        journal_data.append({
                            'Year': year,
                            'Type': 'Journal',
                            'Title': paper['Title'],
                            'Citation Link': paper['Citation Link'],
                            'Venue': paper.get('journal_or_conference', 'N/A'),
                            'Publisher': paper['publisher'],
                            'Cited By': paper['Cited By']
                        })
            journal_df = pd.DataFrame(journal_data)
            if journal_df.empty:
                journal_df = pd.DataFrame(columns=['Year', 'Type', 'Title', 'Citation Link', 'Venue', 'Publisher', 'Cited By'])
            journal_df.to_excel(writer, sheet_name='Journals', index=False)

        if selected_type in ['all', 'conference']:
            conference_data = []
            if conferences:
                for year, papers in sorted(conferences.items()):
                    for paper in papers:
                        conference_data.append({
                            'Year': year,
                            'Type': 'Conference',
                            'Title': paper['Title'],
                            'Citation Link': paper['Citation Link'],
                            'Venue': paper.get('journal_or_conference', 'N/A'),
                            'Publisher': paper['publisher'],
                            'Cited By': paper['Cited By']
                        })
            conference_df = pd.DataFrame(conference_data)
            if conference_df.empty:
                conference_df = pd.DataFrame(columns=['Year', 'Type', 'Title', 'Citation Link', 'Venue', 'Publisher', 'Cited By'])
            conference_df.to_excel(writer, sheet_name='Conferences', index=False)

        if selected_type in ['all', 'miscellaneous']:
            miscellaneous_df = pd.DataFrame(miscellaneous)
            if miscellaneous_df.empty:
                miscellaneous_df = pd.DataFrame(columns=['Title', 'Citation Link', 'Venue', 'Publisher', 'Cited By', 'Year'])
            miscellaneous_df.to_excel(writer, sheet_name='Miscellaneous', index=False)

    output.seek(0)
    return output


def parse_excel(file):
    global authors_list
    df = pd.read_excel(file)
    authors_list = df.iloc[:, 0].tolist()

def save_to_docx(journals, conferences, miscellaneous):
    doc = Document()
    doc.add_heading('Publications Report', level=1)

    def add_table(title, data):
        if data:
            doc.add_heading(title, level=2)
            table = doc.add_table(rows=1, cols=len(data[0]))
            hdr_cells = table.rows[0].cells
            headers = list(data[0].keys())
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            for row in data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(row[header])

    journal_data = []
    for year, papers in sorted(journals.items()):
        for paper in papers:
            journal_data.append({
                'Year': year,
                'Type': 'Journal',
                'Title': paper['Title'],
                'Citation Link': paper['Citation Link'],
                'Venue': paper.get('journal_or_conference', 'N/A'),
                'Publisher': paper['publisher'],
                'Cited By': paper['Cited By']
            })
    add_table('Journals', journal_data)

    conference_data = []
    for year, papers in sorted(conferences.items()):
        for paper in papers:
            conference_data.append({
                'Year': year,
                'Type': 'Conference',
                'Title': paper['Title'],
                'Citation Link': paper['Citation Link'],
                'Venue': paper.get('journal_or_conference', 'N/A'),
                'Publisher': paper['publisher'],
                'Cited By': paper['Cited By']
            })
    add_table('Conferences', conference_data)
    misc_data = []
    for paper in miscellaneous:
        misc_data.append({
            'Title': paper['Title'],
            'Citation Link': paper['Citation Link'],
            'Venue': paper['Venue'],
            'Publisher': paper['publisher'],
            'Cited By': paper['Cited By'],
            'Year': paper['Year']
        })
    add_table('Miscellaneous', misc_data)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files.get('file')
        parse_excel(file)
        global authors_data_for_dashboard, authors_list
        authors_data_for_dashboard = fetch_author_details(authors_list)
        return render_template('dashboard.html', authors=authors_data_for_dashboard)

    return render_template('index.html')

@app.route('/dashboard', methods=['GET', 'POST'])
def dashboard():
    if request.method == 'POST':
        author_name = request.form.get('author_name')
        start_year = request.form.get('start_year')
        end_year = request.form.get('end_year')
        start_year = int(start_year) if start_year and start_year.isdigit() else None
        end_year = int(end_year) if end_year and end_year.isdigit() else None

        journals, conferences, miscellaneous, summary_text = fetch_publications(author_name, start_year, end_year)
        global jour, conf, misc
        jour = journals
        conf = conferences
        misc = miscellaneous
        return render_template('results.html',
                               journals=jour,
                               conferences=conf,
                               miscellaneous=misc,
                               download_url='/download?author_name=' + author_name + '&start_year=' + (str(start_year) if start_year else '') + '&end_year=' + (str(end_year) if end_year else ''),
                               author_name=author_name,
                               start_year=start_year,
                               end_year=end_year,
                               summary_text=summary_text)
    global authors_data_for_dashboard
    return render_template('dashboard.html', authors=authors_data_for_dashboard)

@app.route('/download')
def download():
    author_name = request.args.get('author_name')
    global jour, conf, misc
    publication_type = request.args.get('publication_type')
    
    if publication_type == 'journal':
        excel_file = save_to_excel(journals=jour, selected_type='journal')
        return send_file(excel_file, as_attachment=True, download_name='publications_journals.xlsx')
    elif publication_type == 'conference':
        excel_file = save_to_excel(conferences=conf, selected_type='conference')
        return send_file(excel_file, as_attachment=True, download_name='publications_conf.xlsx')
    elif publication_type == 'docjour':
        docx_file = save_to_docx(jour, {}, [])
        return send_file(docx_file, as_attachment=True, download_name='publications_journals.docx')
    elif publication_type == 'docconf':
        docx_file = save_to_docx({}, conf, [])
        return send_file(docx_file, as_attachment=True, download_name='publications_conf.docx')
    elif publication_type == 'docall':
        docx_file = save_to_docx(jour, conf, misc)
        return send_file(docx_file, as_attachment=True, download_name='publications.docx')

    excel_file = save_to_excel(journals=jour, conferences=conf, miscellaneous=misc, selected_type='all')
    return send_file(excel_file, as_attachment=True, download_name='publications.xlsx')


if __name__ == '__main__':
    app.run(debug=True)
